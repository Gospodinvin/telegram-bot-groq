# docx_export.py — экспорт в DOCX с поддержкой таблиц и графиков
import io
import re
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Any

from docx import Document
from docx.shared import Pt, Mm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from gost_standards import get_format_info, WORK_TYPES
from utils.diploma_utils import parse_diploma_text

def _add_image_to_paragraph(paragraph, image_bytes: bytes, width: Optional[Mm] = None):
    from docx.shared import Inches
    run = paragraph.add_run()
    if width:
        run.add_picture(io.BytesIO(image_bytes), width=width)
    else:
        run.add_picture(io.BytesIO(image_bytes), width=Inches(5.5))

def _set_run_font(run, font_name: str = 'Times New Roman', size: int = 14, bold: bool = False):
    font = run.font
    font.name = font_name
    font.size = Pt(size)
    font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def _set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                          first_line_indent=Mm(12.5),
                          line_spacing=1.5,
                          space_after=Pt(0),
                          space_before=Pt(0)):
    pf = p.paragraph_format
    pf.alignment = alignment
    pf.first_line_indent = first_line_indent
    pf.space_after = space_after
    pf.space_before = space_before
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line_spacing

def _add_page_number_field(paragraph):
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

def _build_title_page(doc, styles, title: str, meta: dict, work_type: str):
    wt = WORK_TYPES.get(work_type, WORK_TYPES['diploma'])
    s_meta = styles['Meta']
    s_title = styles['Title']

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(meta.get('university', 'ФЕДЕРАЛЬНЫЙ УНИВЕРСИТЕТ').upper())
    _set_run_font(r, s_meta.font.name, 14, bold=False)

    if meta.get('faculty'):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(meta['faculty'])
        _set_run_font(r, s_meta.font.name, 14)

    if meta.get('department'):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("Кафедра " + meta['department'])
        _set_run_font(r, s_meta.font.name, 14)

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(wt['name'].upper())
    _set_run_font(r, s_title.font.name, 16, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("на тему:")
    _set_run_font(r, s_meta.font.name, 14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    _set_run_font(r, s_title.font.name, 16, bold=True)

    doc.add_paragraph()
    doc.add_paragraph()

    author = meta.get('author', 'Иванов И. И.')
    group = meta.get('group', '')
    supervisor = meta.get('supervisor', 'Петров П. П.')
    supervisor_title = meta.get('supervisor_title', 'доцент, к.т.н.')

    table = doc.add_table(rows=5, cols=2)
    table.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    table.autofit = False
    table.allow_autofit = False
    table.columns[0].width = Mm(60)
    table.columns[1].width = Mm(70)

    cells = [
        ("Выполнил:", author),
        ("", group),
        ("", ""),
        ("Научный руководитель:", supervisor),
        ("", supervisor_title),
    ]
    for i, (left, right) in enumerate(cells):
        row = table.rows[i]
        row.cells[0].text = left
        row.cells[1].text = right
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    _set_run_font(run, s_meta.font.name, 14)

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    city = meta.get('city', 'Екатеринбург')
    year = meta.get('year', datetime.now().year)
    r = p.add_run(city + " — " + str(year))
    _set_run_font(r, s_meta.font.name, 14)

def _build_abstract(doc, styles, meta: dict, work_type: str):
    wt = WORK_TYPES.get(work_type, WORK_TYPES['diploma'])
    s_body = styles['Body']
    s_h1 = styles['Heading1']

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("РЕФЕРАТ")
    _set_run_font(r, s_h1.font.name, 14, bold=True)

    volume = meta.get('volume_pages', wt['min_pages'])
    chapters = meta.get('chapters_count', wt['chapters'])
    figures = meta.get('figures_count', 0)
    tables = meta.get('tables_count', 0)
    sources = meta.get('sources_count', 10)
    appendices = meta.get('appendices_count', 0)

    volume_line = (
        f"Пояснительная записка содержит {volume} стр., "
        f"{chapters} {'главы' if chapters < 5 else 'глав'}, "
        f"{figures} рис., {tables} табл., {sources} ист."
    )
    if appendices:
        volume_line += f", {appendices} прил."

    p = doc.add_paragraph()
    r = p.add_run(volume_line)
    _set_run_font(r, s_body.font.name, 14)
    _set_paragraph_format(p, WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_indent=Mm(12.5))

    fields = [
        ("Ключевые слова", meta.get('keywords', '')),
        ("Объект исследования", meta.get('object', '')),
        ("Предмет исследования", meta.get('subject', '')),
        ("Цель работы", meta.get('goal', '')),
        ("Методы исследования", meta.get('methods', '')),
        ("Результаты", meta.get('results', '')),
        ("Область применения", meta.get('application', '')),
    ]
    for label, value in fields:
        if value:
            p = doc.add_paragraph()
            r = p.add_run(label + ": ")
            _set_run_font(r, s_body.font.name, 14, bold=True)
            r = p.add_run(value)
            _set_run_font(r, s_body.font.name, 14)
            _set_paragraph_format(p, WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_indent=Mm(12.5))

def _build_toc(doc, styles, sections: list):
    s_text = styles['Body']
    s_h1 = styles['Heading1']

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("ОГЛАВЛЕНИЕ")
    _set_run_font(r, s_h1.font.name, 14, bold=True)

    entries = []
    page = 3

    entries.append(("ВВЕДЕНИЕ", "3"))

    for heading, _ in sections:
        if not heading:
            continue
        h_up = heading.upper()
        if any(x in h_up for x in ['ВВЕДЕНИЕ', 'ЗАКЛЮЧЕНИЕ', 'СПИСОК', 'ПРИЛОЖЕНИЕ']):
            continue
        if re.match(r'^глава\s+\d+', heading, re.I):
            heading = "ГЛАВА " + re.search(r'\d+', heading).group()
        entries.append((heading, str(page)))
        page += 3

    entries.append(("ЗАКЛЮЧЕНИЕ", str(page)))
    entries.append(("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ", str(page + 2)))

    for text, pg in entries:
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.tab_stops.add_tab_stop(Inches(6.1), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)

        r = p.add_run(text)
        _set_run_font(r, s_text.font.name, 14)
        p.add_run("\t")
        r = p.add_run(pg)
        _set_run_font(r, s_text.font.name, 14)

def build_gost_docx(
    title: str,
    sections: List[Tuple[str, str]],
    standard_key: str = "gost_7.32-2017",
    meta: dict = None,
    work_type: str = "diploma",
    images: Optional[Dict[str, bytes]] = None
) -> io.BytesIO:
    fmt = get_format_info(standard_key)
    font_name = fmt.get('font', 'Times New Roman')
    font_size = fmt.get('font_size', 14)
    line_sp = fmt.get('line_spacing', 1.5)
    margins = fmt.get('margins', {'left': 30, 'right': 15, 'top': 20, 'bottom': 20})

    doc = Document()

    section = doc.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.top_margin = Mm(margins['top'])
    section.bottom_margin = Mm(margins['bottom'])
    section.left_margin = Mm(margins['left'])
    section.right_margin = Mm(margins['right'])

    style = doc.styles['Normal']
    style.font.name = font_name
    style.font.size = Pt(font_size)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

    styles = {
        'Title': style,
        'Heading1': style,
        'Body': style,
        'Meta': style,
    }

    if meta:
        _build_title_page(doc, styles, title, meta, work_type)
        doc.add_section(start_type=2)

    has_abstract = meta and meta.get('include_abstract', True)
    if has_abstract:
        _build_abstract(doc, styles, meta, work_type)
        doc.add_page_break()

    _build_toc(doc, styles, sections)
    doc.add_page_break()

    images = images or {}
    for heading, body in sections:
        if heading:
            p = doc.add_paragraph()
            h_up = heading.upper()
            is_major = heading.isupper() or any(x in h_up for x in [
                'ВВЕДЕНИЕ', 'ЗАКЛЮЧЕНИЕ', 'СПИСОК', 'ЛИТЕРАТУР', 'ПРИЛОЖЕНИЕ'
            ])
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if is_major else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(heading)
            _set_run_font(r, font_name, font_size, bold=True)
            _set_paragraph_format(p, p.alignment, first_line_indent=Mm(0), line_spacing=line_sp, space_after=Pt(12))

        lines = body.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue

            table_marker = re.search(r'\[ТАБЛИЦА\s*(\d+)\]', line)
            figure_marker = re.search(r'\[(РИСУНОК|ГРАФИК)\s*(\d+)\]', line)

            if table_marker:
                table_id = table_marker.group(1)
                key_candidates = [f"table_{table_id}", f"таблица_{table_id}", f"table{table_id}"]
                img_bytes = None
                for key in key_candidates:
                    if key in images:
                        img_bytes = images[key]
                        break
                if img_bytes:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    _add_image_to_paragraph(p, img_bytes, width=Mm(160))
                else:
                    p = doc.add_paragraph()
                    r = p.add_run(f"[ТАБЛИЦА {table_id} — изображение не найдено]")
                    _set_run_font(r, font_name, font_size)
                continue

            elif figure_marker:
                fig_type = figure_marker.group(1)
                fig_id = figure_marker.group(2)
                key_candidates = [f"figure_{fig_id}", f"рисунок_{fig_id}", f"chart_{fig_id}", f"fig{fig_id}"]
                img_bytes = None
                for key in key_candidates:
                    if key in images:
                        img_bytes = images[key]
                        break
                if img_bytes:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    _add_image_to_paragraph(p, img_bytes, width=Mm(160))
                else:
                    p = doc.add_paragraph()
                    r = p.add_run(f"[{fig_type} {fig_id} — изображение не найдено]")
                    _set_run_font(r, font_name, font_size)
                continue

            p = doc.add_paragraph()
            if re.match(r'^\d+\.', line) or re.match(r'^\[\d+\]', line):
                p.paragraph_format.left_indent = Mm(10)
                p.paragraph_format.first_line_indent = Mm(-10)
            r = p.add_run(line)
            _set_run_font(r, font_name, font_size)
            _set_paragraph_format(p, WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_indent=Mm(12.5), line_spacing=line_sp)

    for i, sect in enumerate(doc.sections):
        if i == 0:
            sect.footer.is_linked_to_previous = False
            for p in sect.footer.paragraphs:
                p.clear()
            continue

        sect.footer.is_linked_to_previous = False
        footer = sect.footer
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_page_number_field(footer_para)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

def build_simple_docx(title: str, body: str) -> io.BytesIO:
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(14)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    _set_run_font(r, 'Times New Roman', 16, bold=True)

    p = doc.add_paragraph()
    r = p.add_run("Дата: " + datetime.now().strftime('%d.%m.%Y %H:%M'))
    _set_run_font(r, 'Times New Roman', 12)

    for line in body.split('\n'):
        line = line.strip()
        if not line:
            continue
        p = doc.add_paragraph()
        r = p.add_run(line)
        _set_run_font(r, 'Times New Roman', 14)
        _set_paragraph_format(p, WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=Mm(0), line_spacing=1.5)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf